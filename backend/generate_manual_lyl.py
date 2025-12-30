"""
系统使用说明书生成脚本_lyl
使用docx库生成Word文档
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
import os


def create_manual_lyl():
    """创建系统使用说明书_lyl"""
    doc = Document()
    
    # 设置文档标题样式_lyl
    title = doc.add_heading('个性化知识问答系统', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('系统使用说明书')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0] if subtitle.runs else subtitle.add_run()
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.bold = True
    
    # 添加版本和日期信息_lyl
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(f'版本: 1.0.0\n')
    info.add_run(f'日期: {datetime.now().strftime("%Y年%m月%d日")}')
    
    doc.add_paragraph()  # 空行
    
    # 目录_lyl
    doc.add_heading('目录', level=1)
    toc_items = [
        '一、项目概述',
        '二、系统功能介绍',
        '三、技术架构',
        '四、安装部署说明',
        '五、exe独立运行说明',
        '六、使用教程',
        '七、常见问题',
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Number')
    
    doc.add_page_break()
    
    # 一、项目概述_lyl
    doc.add_heading('一、项目概述', level=1)
    
    doc.add_heading('1.1 项目背景', level=2)
    doc.add_paragraph(
        '随着人工智能技术的快速发展，大语言模型（LLM）在自然语言处理领域展现出了强大的能力。'
        '然而，通用大模型往往缺乏对特定领域知识的深入理解。本项目旨在构建一个"个性化知识问答系统"，'
        '通过RAG（检索增强生成）技术，将用户的私有知识库与大语言模型相结合，'
        '实现基于个人知识的智能问答功能。'
    )
    
    doc.add_heading('1.2 项目目标', level=2)
    objectives = [
        '提供友好的用户界面，支持多种文档格式的上传和管理',
        '实现基于向量检索的知识库搜索功能',
        '集成大语言模型API，提供智能对话能力',
        '支持流式输出，提升用户交互体验',
        '提供对话历史管理功能，支持历史对话的查看和继续',
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')
    
    doc.add_heading('1.3 技术特点', level=2)
    features = [
        'LangChain Agent模式：将RAG作为工具绑定到Agent，实现智能决策',
        'SSE流式输出：实时显示AI回复，提升交互体验',
        '引用溯源：回复内容带有引用标记，可追溯知识来源',
        '工具调用可视化：展示AI调用知识库检索的过程',
    ]
    for feat in features:
        doc.add_paragraph(feat, style='List Bullet')
    
    # 二、系统功能介绍_lyl
    doc.add_heading('二、系统功能介绍', level=1)
    
    doc.add_heading('2.1 智能对话', level=2)
    doc.add_paragraph(
        '系统的核心功能是智能对话。用户可以输入问题，系统会自动判断是否需要检索知识库，'
        '并结合检索到的相关内容生成回答。回答中会包含引用标记（如[1][2]），'
        '点击可以查看对应的知识来源。'
    )
    dialog_features = [
        '支持开启/关闭知识库检索',
        '流式输出，实时显示回复内容',
        '工具调用卡片展示检索过程',
        '引用标记可点击查看来源',
        '支持新建对话和继续历史对话',
    ]
    for feat in dialog_features:
        doc.add_paragraph(feat, style='List Bullet')
    
    doc.add_heading('2.2 知识库管理', level=2)
    doc.add_paragraph(
        '用户可以上传多种格式的文档到知识库，系统会自动进行文本分割和向量化处理，'
        '存储到FAISS向量数据库中，供后续检索使用。'
    )
    doc.add_paragraph('支持的文档格式：')
    formats = ['TXT - 纯文本文件', 'PDF - PDF文档', 'DOCX/DOC - Word文档', 'MD - Markdown文件']
    for fmt in formats:
        doc.add_paragraph(fmt, style='List Bullet')
    
    doc.add_heading('2.3 对话历史', level=2)
    doc.add_paragraph(
        '系统会自动保存每次对话的记录，用户可以在对话历史页面查看所有历史对话，'
        '支持搜索、删除功能，还可以点击"继续对话"按钮恢复历史对话并继续交流。'
    )
    
    # 三、技术架构_lyl
    doc.add_heading('三、技术架构', level=1)
    
    doc.add_heading('3.1 整体架构', level=2)
    doc.add_paragraph('本系统采用前后端分离架构：')
    arch_items = [
        '前端：React + TypeScript + Vite + Ant Design',
        '后端：Python + FastAPI + LangChain',
        '数据库：SQLite (aiosqlite异步驱动)',
        '向量存储：FAISS',
        '大语言模型：DeepSeek API',
        '嵌入模型：BGE-M3 (SiliconFlow API)',
    ]
    for item in arch_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('3.2 核心模块', level=2)
    
    modules = [
        ('AgentService', '基于LangChain的智能代理服务，负责对话处理和工具调用'),
        ('RAG Tool', 'RAG检索工具，封装向量搜索功能，作为Agent的工具使用'),
        ('VectorStoreService', '向量存储服务，管理FAISS向量数据库的读写'),
        ('DocumentService', '文档处理服务，负责文档加载、分割和存储'),
        ('ConversationService', '对话管理服务，负责对话和消息的CRUD操作'),
    ]
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '模块名称'
    hdr_cells[1].text = '功能描述'
    for name, desc in modules:
        row_cells = table.add_row().cells
        row_cells[0].text = name
        row_cells[1].text = desc
    
    doc.add_paragraph()
    
    doc.add_heading('3.3 数据流程', level=2)
    doc.add_paragraph('智能对话的数据流程如下：')
    flow_steps = [
        '用户输入问题',
        '前端通过SSE请求发送到后端',
        'Agent判断是否需要检索知识库',
        '如需检索，调用RAG工具搜索向量库',
        '将检索结果注入到对话上下文',
        '调用DeepSeek API生成回答',
        '流式返回回答内容给前端',
        '保存对话记录到数据库',
    ]
    for i, step in enumerate(flow_steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    # 四、安装部署说明_lyl
    doc.add_heading('四、安装部署说明', level=1)

    doc.add_heading('4.1 环境要求', level=2)
    doc.add_paragraph('方式一（本地开发）：')
    requirements = [
        'Python 3.10+',
        'Node.js 18+',
        'uv (Python包管理器)',
        'npm (Node包管理器)',
    ]
    for req in requirements:
        doc.add_paragraph(req, style='List Bullet')

    doc.add_paragraph('方式二（Docker部署）：')
    docker_reqs = [
        'Docker Desktop',
        'docker-compose',
    ]
    for req in docker_reqs:
        doc.add_paragraph(req, style='List Bullet')

    doc.add_heading('4.2 配置环境变量', level=2)
    doc.add_paragraph('在项目根目录创建 .env 文件，配置以下内容：')

    env_config = '''LLM_API_KEY=你的DeepSeek API Key
LLM_BASE_URL=https://api.deepseek.com/v1
LLM_MODEL=deepseek-chat
EMBEDDING_MODEL_API_KEY=你的SiliconFlow API Key
EMBEDDING_MODEL_BASE_URL=https://api.siliconflow.cn/v1
EMBEDDING_MODEL=BAAI/bge-m3'''

    p = doc.add_paragraph()
    run = p.add_run(env_config)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    doc.add_heading('4.3 方式一：本地开发部署', level=2)

    doc.add_paragraph('1. 安装后端依赖')
    p = doc.add_paragraph()
    run = p.add_run('cd backend && uv sync')
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    doc.add_paragraph('2. 安装前端依赖')
    p = doc.add_paragraph()
    run = p.add_run('cd frontend && npm install')
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    doc.add_paragraph('3. 启动后端服务')
    p = doc.add_paragraph()
    run = p.add_run('uv run uvicorn backend.app.main:app --reload --port 8000')
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    doc.add_paragraph('4. 启动前端服务（新终端）')
    p = doc.add_paragraph()
    run = p.add_run('cd frontend && npm run dev')
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    doc.add_paragraph('5. 访问系统：打开浏览器访问 http://localhost:5173')

    doc.add_heading('4.4 方式二：Docker 部署（推荐）', level=2)
    doc.add_paragraph('使用 Docker 可以一键部署整个系统，无需手动安装依赖。')

    doc.add_paragraph('1. 确保已安装 Docker Desktop 并启动')

    doc.add_paragraph('2. 在项目根目录执行构建和启动命令')
    p = doc.add_paragraph()
    run = p.add_run('docker-compose up -d --build')
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    doc.add_paragraph('3. 访问系统')
    docker_access = [
        '前端界面：http://localhost',
        '后端API：http://localhost:8000',
        'API文档：http://localhost:8000/docs',
    ]
    for item in docker_access:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph('4. 查看日志')
    p = doc.add_paragraph()
    run = p.add_run('docker-compose logs -f')
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    doc.add_paragraph('5. 停止服务')
    p = doc.add_paragraph()
    run = p.add_run('docker-compose down')
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    doc.add_paragraph('注意：数据会持久化到 ./data 目录，包括数据库、向量存储和上传的文档。')

    # 五、exe独立运行说明_lyl
    doc.add_heading('五、exe独立运行说明', level=1)

    doc.add_heading('5.1 简介', level=2)
    doc.add_paragraph(
        '本系统支持打包成单个可执行文件（exe），无需安装Python环境或其他依赖，'
        '双击即可运行。适合不熟悉技术的用户或需要快速部署的场景。'
    )

    doc.add_heading('5.2 运行exe文件', level=2)
    exe_steps = [
        '确保已获取到"知识问答系统.exe"文件',
        '将exe文件放到一个独立的文件夹中（首次运行会创建data目录）',
        '在同目录下创建.env文件，配置API密钥（参见4.2节）',
        '双击运行exe文件',
        '等待系统启动，浏览器会自动打开 http://localhost:8000',
        '使用完毕后，在命令行窗口按 Ctrl+C 停止服务',
    ]
    for i, step in enumerate(exe_steps, 1):
        doc.add_paragraph(f'{i}. {step}')

    doc.add_heading('5.3 exe文件目录结构', level=2)
    doc.add_paragraph('运行后会自动创建以下目录结构：')
    dir_structure = '''知识问答系统.exe
.env                    # 环境变量配置文件（需手动创建）
data/
  ├── knowledge_qa.db   # SQLite数据库
  ├── vector_store/     # FAISS向量存储
  └── documents/        # 上传的文档'''
    p = doc.add_paragraph()
    run = p.add_run(dir_structure)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    doc.add_heading('5.4 打包exe文件（开发者）', level=2)
    doc.add_paragraph('如需自行打包exe文件，执行以下命令：')
    p = doc.add_paragraph()
    run = p.add_run('uv run python build_exe_lyl.py')
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    doc.add_paragraph('打包完成后，exe文件位于 dist/知识问答系统.exe')

    # 六、使用教程_lyl
    doc.add_heading('六、使用教程', level=1)
    
    doc.add_heading('6.1 上传知识库文档', level=2)
    upload_steps = [
        '点击左侧菜单"知识库管理"',
        '点击"上传文档"按钮',
        '选择要上传的文档文件（支持txt/pdf/docx/md格式）',
        '等待上传和处理完成',
        '在文档列表中查看已上传的文档',
    ]
    for i, step in enumerate(upload_steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    doc.add_heading('6.2 开始智能对话', level=2)
    chat_steps = [
        '点击左侧菜单"智能对话"',
        '在底部输入框输入您的问题',
        '确保"使用知识库"开关已开启（如需检索知识库）',
        '点击"发送"按钮或按Enter键发送',
        '等待AI回复，可以看到流式输出的内容',
        '点击回复中的引用标记[1][2]等查看知识来源',
    ]
    for i, step in enumerate(chat_steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    doc.add_heading('6.3 查看和继续历史对话', level=2)
    history_steps = [
        '点击左侧菜单"对话历史"',
        '在列表中可以看到所有历史对话',
        '使用搜索框可以按标题搜索对话',
        '点击"继续对话"按钮可以恢复该对话并继续交流',
        '点击"删除"按钮可以删除不需要的对话记录',
    ]
    for i, step in enumerate(history_steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    # 七、常见问题_lyl
    doc.add_heading('七、常见问题', level=1)

    faqs = [
        ('Q: 上传文档后为什么搜索不到内容？',
         'A: 请确保文档内容不为空，且等待向量化处理完成。可以在知识库管理页面查看文档的分块数量。'),
        ('Q: 对话时提示"知识库为空"怎么办？',
         'A: 请先上传至少一个文档到知识库，系统需要有知识内容才能进行检索。'),
        ('Q: 如何关闭知识库检索，直接和AI对话？',
         'A: 在对话输入框上方有"使用知识库"开关，关闭即可直接与AI对话而不检索知识库。'),
        ('Q: 系统支持哪些大语言模型？',
         'A: 目前系统使用DeepSeek API，理论上支持任何兼容OpenAI API格式的模型服务。'),
        ('Q: 运行exe文件报错"DLL load failed"怎么办？',
         'A: 这通常是因为系统缺少Visual C++运行库。请下载安装Microsoft Visual C++ Redistributable最新版本。'),
        ('Q: exe运行时提示找不到.env文件怎么办？',
         'A: 请在exe文件同目录下创建.env文件，并配置好API密钥（参见安装部署说明章节）。'),
        ('Q: 数据保存在哪里？',
         'A: 所有数据（数据库、向量存储、上传的文档）都保存在exe同目录的data文件夹中，可以备份此文件夹。'),
    ]
    
    for q, a in faqs:
        doc.add_paragraph(q).runs[0].bold = True
        doc.add_paragraph(a)
        doc.add_paragraph()  # 空行
    
    # 保存文档_lyl
    output_path = os.path.join(os.path.dirname(__file__), '..', '系统使用说明书.docx')
    doc.save(output_path)
    print(f'使用说明书已生成: {os.path.abspath(output_path)}')
    return output_path


if __name__ == '__main__':
    create_manual_lyl()

